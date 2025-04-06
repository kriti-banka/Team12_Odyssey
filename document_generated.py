import os
import csv
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Remove direct import from app.py to avoid circular imports
# Instead, we'll use function parameters to pass RAG functionality

def load_company_data(csv_path="data.csv"):
    """Load company data from CSV file into a dictionary."""
    company_data = {}
    try:
        with open(csv_path, 'r', encoding='utf-8') as file:
            reader = csv.reader(file)
            next(reader)  # Skip header row
            for row in reader:
                if len(row) >= 2:
                    key, value = row[0], row[1]
                    company_data[key] = value
    except Exception as e:
        print(f"Error loading CSV data: {e}")
    
    return company_data

def add_page_number(paragraph):
    """Add page numbers to the document footer."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)
    
    return paragraph

def apply_document_styling(doc):
    """Apply consistent styling to the document."""
    # Add document styles
    styles = doc.styles
    
    # Heading 1 style
    if 'Heading 1' in styles:
        style = styles['Heading 1']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(24)
        font.bold = True
        font.color.rgb = RGBColor(0, 77, 113)  # Dark blue
    
    # Heading 2 style
    if 'Heading 2' in styles:
        style = styles['Heading 2']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(14)
        font.bold = True
        font.color.rgb = RGBColor(0, 77, 113)  # Dark blue
    
    # Normal text style
    if 'Normal' in styles:
        style = styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        style.paragraph_format.space_after = Pt(6)
    
    # List Bullet style
    if 'List Bullet' in styles:
        style = styles['List Bullet']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
    
    # List Number style
    if 'List Number' in styles:
        style = styles['List Number']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(12)
    
    # Add footer with page numbers
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    add_page_number(paragraph)
    
    return doc

def add_cover_page(doc, company_data):
    """Add a professional cover page to the document following Page 1 specs."""
    # Add title - Heading 1, center-aligned, font size 24pt, bold
    title = doc.add_paragraph(style='Heading 1')
    title_run = title.add_run("PROPOSAL DOCUMENT")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add company name with same styling
    company_name = doc.add_paragraph(style='Heading 1')
    company_name_run = company_name.add_run(company_data.get('Company Legal Name', 'Company Name'))
    company_name.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company_name.space_after = Pt(36)
    
    # Prepared For/By - Heading 2, font size 14pt, bold
    prepared_by = doc.add_paragraph(style='Heading 2')
    prepared_by_run = prepared_by.add_run("Prepared By:")
    prepared_by.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    company = doc.add_paragraph()
    company_run = company.add_run(company_data.get('Company Legal Name', 'Company Name'))
    company_run.bold = True
    company.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    company.space_after = Pt(24)
    
    # Submission Date - Normal, font size 12pt
    from datetime import datetime
    date_para = doc.add_paragraph()
    date_para.add_run("Submission Date: ").bold = True
    date_para.add_run(datetime.now().strftime("%B %d, %Y"))
    date_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    date_para.space_after = Pt(24)
    
    # Contact Info - Using separate Paragraphs, center-aligned
    contact_info = doc.add_paragraph()
    contact_info.add_run("Contact Information:").bold = True
    contact_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    address = doc.add_paragraph()
    address.add_run(company_data.get('Principal Business Address', ''))
    address.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    phone = doc.add_paragraph()
    phone.add_run(f"Phone: {company_data.get('Phone Number', '')}")
    phone.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    email = doc.add_paragraph()
    email.add_run(f"Email: {company_data.get('Email Address', '')}")
    email.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    email.space_after = Pt(36)
    
    # Authorized Representative - Normal, bold name, title italicized
    auth_rep = doc.add_paragraph()
    auth_rep.add_run("Authorized Representative: ").bold = True
    auth_rep.add_run(company_data.get('Authorized Representative', '')).bold = True
    auth_rep.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    auth_title = doc.add_paragraph()
    title_run = auth_title.add_run(company_data.get('Authorized Representative Title', ''))
    title_run.italic = True
    auth_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add page break at the end of cover page
    doc.add_page_break()
    
    return doc

def add_executive_summary(doc, company_data, rag_content=None):
    """Add executive summary page following Page 2 specs."""
    # Section Title - Heading 1, font size 18pt, underline
    exec_summary = doc.add_paragraph(style='Heading 1')
    exec_summary_run = exec_summary.add_run("EXECUTIVE SUMMARY")
    exec_summary_run.underline = True
    exec_summary_run.font.size = Pt(18)
    
    # Body Text - Normal, 12pt, left-aligned
    # Use provided RAG content if available, otherwise use generic content
    if rag_content and rag_content.get('executive_summary'):
        summary_content = rag_content['executive_summary']
    else:
        # Fallback content if RAG content not available
        summary_content = f"{company_data.get('Company Legal Name', 'Our company')} is pleased to submit this proposal for providing {company_data.get('Services Provided', 'professional services')}. With {company_data.get('Years of Experience in Temporary Staffing', '')} of experience, we are uniquely qualified to deliver high-quality solutions that meet your needs."
    
    doc.add_paragraph(summary_content)
    
    # Key Benefits section - Bold to emphasize key phrases
    benefits_title = doc.add_paragraph()
    benefits_title.add_run("KEY BENEFITS").bold = True
    
    # Bullet Points - Use Paragraph().style = 'List Bullet'
    benefits = [
        "Experienced team with specialized expertise",
        f"{company_data.get('Years of Experience in Temporary Staffing', '')} of industry experience",
        "Proven track record of success",
        "Customized solutions tailored to your needs",
        "Exceptional customer service and support"
    ]
    
    for benefit in benefits:
        bullet = doc.add_paragraph(style='List Bullet')
        parts = benefit.split(' with ')
        if len(parts) > 1:
            bullet.add_run(parts[0]).bold = True
            bullet.add_run(f" with {parts[1]}")
        else:
            bullet.add_run(benefit)
    
    # Add page break at the end
    doc.add_page_break()
    
    return doc

def add_company_overview(doc, company_data):
    """Add company overview page following Page 3 specs."""
    # Section Title - Heading 1
    doc.add_paragraph("COMPANY OVERVIEW & EXPERIENCE", style='Heading 1')
    
    # Subsections - Use bold Normal
    company_info = doc.add_paragraph()
    company_info.add_run(f"{company_data.get('Company Legal Name', 'Our company')} is a {company_data.get('Business Structure', 'professional organization')} established {company_data.get('Company Length of Existence', '')} ago. We specialize in providing {company_data.get('Services Provided', 'staffing services')}.")
    
    # Company identifiers and registrations - Bold
    identifiers = doc.add_paragraph()
    
    # DUNS, NAICS, CAGE Code
    if 'DUNS Number' in company_data:
        identifiers.add_run("DUNS Number: ").bold = True
        identifiers.add_run(f"{company_data.get('DUNS Number', '')}\n")
    
    if 'CAGE Code' in company_data:
        identifiers.add_run("CAGE Code: ").bold = True
        identifiers.add_run(f"{company_data.get('CAGE Code', '')}\n")
    
    if 'NAICS Codes' in company_data:
        identifiers.add_run("NAICS Codes: ").bold = True
        identifiers.add_run(f"{company_data.get('NAICS Codes', '')}\n")
    
    if 'State Registration Number' in company_data:
        identifiers.add_run("State Registration: ").bold = True
        identifiers.add_run(f"{company_data.get('State Registration Number', '')}")
    
    # Past Experience - Use List Bullet
    doc.add_paragraph("PAST EXPERIENCE", style='Heading 2')
    
    experience_intro = doc.add_paragraph()
    experience_intro.add_run(f"With {company_data.get('Years of Experience in Temporary Staffing', '')} in the industry, we have successfully delivered staffing solutions across various sectors. Our experience includes:")
    
    # Sample experience points (would ideally come from RAG or additional data)
    experiences = [
        "Government agency staffing and workforce management",
        "Corporate staffing for Fortune 500 companies",
        "Healthcare and medical staffing solutions",
        "IT and technical professional placements",
        "Administrative and support staff services"
    ]
    
    for exp in experiences:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.add_run(exp)
    
    # Special Qualifications - Italics or bold keywords
    doc.add_paragraph("SPECIAL QUALIFICATIONS", style='Heading 2')
    
    qualifications = [
        "Rapid deployment of qualified personnel",
        "Stringent vetting and security clearance processes",
        "Continuous performance monitoring",
        "Dedicated account management",
        f"{company_data.get('Historically Underutilized Business/DBE Status', 'Business status')}"
    ]
    
    for qual in qualifications:
        bullet = doc.add_paragraph(style='List Bullet')
        if ":" in qual:
            parts = qual.split(":")
            bullet.add_run(f"{parts[0]}:").bold = True
            bullet.add_run(parts[1])
        else:
            bullet.add_run(qual)
    
    # Add page break at the end
    doc.add_page_break()
    
    return doc

def add_scope_and_approach(doc, rag_content=None):
    """Add scope of work page following Page 4 specs."""
    # Section Title - Heading 1
    doc.add_paragraph("SCOPE OF WORK & SOLUTION APPROACH", style='Heading 1')
    
    # Use provided RAG content if available, otherwise use generic content
    if rag_content and rag_content.get('approach'):
        approach_content = rag_content['approach']
    else:
        # Fallback content if RAG content not available
        approach_content = "Our comprehensive approach to service delivery ensures high-quality results through systematic planning, execution, and monitoring processes."
    
    doc.add_paragraph(approach_content)
    
    # Subsection Titles - Heading 2
    doc.add_paragraph("IMPLEMENTATION METHODOLOGY", style='Heading 2')
    
    # Step-by-Step Plan - Use numbered list
    doc.add_paragraph("Our implementation follows a proven step-by-step process:")
    
    steps = [
        "Requirements gathering and analysis",
        "Staffing plan development",
        "Candidate identification and screening",
        "Selection and onboarding",
        "Performance monitoring and reporting",
        "Continuous improvement"
    ]
    
    for i, step in enumerate(steps, 1):
        bullet = doc.add_paragraph(style='List Number')
        bullet.add_run(f"{step}").bold = True
        # Add a brief description
        bullet.add_run(f" - Ensuring alignment with your specific needs and objectives.")
    
    # Timeline - Use a table
    doc.add_paragraph("IMPLEMENTATION TIMELINE", style='Heading 2')
    doc.add_paragraph("Below is our estimated timeline for full implementation:")
    
    timeline_table = doc.add_table(rows=7, cols=2)
    timeline_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = timeline_table.rows[0].cells
    header_cells[0].text = "Phase"
    header_cells[1].text = "Timeline"
    
    # Style header row
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.bold = True
    
    # Timeline data
    timeline_data = [
        ("Initial Consultation", "Week 1"),
        ("Requirements Analysis", "Week 1-2"),
        ("Candidate Sourcing", "Week 2-3"),
        ("Interviews & Selection", "Week 3-4"),
        ("Onboarding", "Week 4-5"),
        ("Performance Monitoring", "Ongoing")
    ]
    
    for i, (phase, time) in enumerate(timeline_data, 1):
        row = timeline_table.rows[i].cells
        row[0].text = phase
        row[1].text = time
    
    # Tools & Deliverables - Use bold keywords
    doc.add_paragraph("TOOLS & DELIVERABLES", style='Heading 2')
    
    deliverables = [
        "Staffing plan documentation",
        "Candidate selection reports",
        "Performance dashboards",
        "Quality assurance documentation",
        "Regular status reports"
    ]
    
    for deliv in deliverables:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet.add_run(f"✓ {deliv}")
    
    # Add page break at the end
    doc.add_page_break()
    
    return doc

def add_personnel_section(doc, company_data):
    """Add key personnel page following Page 5 specs."""
    # Section Title - Heading 1
    doc.add_paragraph("KEY PERSONNEL & STAFFING PLAN", style='Heading 1')
    
    doc.add_paragraph("Our team brings extensive expertise and experience to ensure the successful delivery of all project requirements.")
    
    # Find all key personnel entries in the company data
    key_personnel = {k: v for k, v in company_data.items() if k.startswith('Key Personnel')}
    
    if key_personnel:
        for role, name in key_personnel.items():
            # Personnel Names - Heading 2, bold
            role_desc = role.replace('Key Personnel â€" ', '').replace('Key Personnel – ', '')
            
            person_heading = doc.add_paragraph(style='Heading 2')
            person_heading.add_run(name)
            
            # Role heading
            role_para = doc.add_paragraph()
            role_para.add_run(f"Role: {role_desc}").bold = True
            
            # Background - Normal with bullet list
            background = doc.add_paragraph()
            background.add_run("Background & Qualifications:").bold = True
            
            # Example qualifications (these would ideally come from another data source)
            qualifications = [
                f"Over 10 years of experience in {role_desc}",
                "Expert in project management and staff coordination",
                "Certified Professional in relevant field",
                "Proven track record of successful project delivery"
            ]
            
            for qual in qualifications:
                bullet = doc.add_paragraph(style='List Bullet')
                bullet.add_run(qual)
    else:
        doc.add_paragraph("Key personnel information not available.")
    
    # Org Chart representation as nested bullets
    doc.add_paragraph("ORGANIZATIONAL STRUCTURE", style='Heading 2')
    
    doc.add_paragraph("Our staffing organization is structured to provide clear lines of communication and accountability:")
    
    # Level 1
    l1 = doc.add_paragraph(style='List Bullet')
    l1.add_run("Project Manager").bold = True
    
    # Level 2 - indented bullets
    l2a = doc.add_paragraph(style='List Bullet')
    l2a.paragraph_format.left_indent = Inches(0.5)
    l2a.add_run("Technical Lead").bold = True
    
    l2b = doc.add_paragraph(style='List Bullet')
    l2b.paragraph_format.left_indent = Inches(0.5)
    l2b.add_run("Administrative Lead").bold = True
    
    # Level 3 - further indented
    l3a = doc.add_paragraph(style='List Bullet')
    l3a.paragraph_format.left_indent = Inches(1.0)
    l3a.add_run("Technical Staff")
    
    l3b = doc.add_paragraph(style='List Bullet')
    l3b.paragraph_format.left_indent = Inches(1.0)
    l3b.add_run("Administrative Support Staff")
    
    # Availability - Emphasize with bold
    avail_para = doc.add_paragraph()
    avail_para.add_run("AVAILABILITY: ").bold = True
    avail_para.add_run("Our team provides ").italic = False
    avail_para.add_run("24/7 availability").bold = True
    avail_para.add_run(" for critical issues with standard support during business hours for routine matters.")
    
    # Add page break at the end
    doc.add_page_break()
    
    return doc

def add_pricing_section(doc, company_data, rag_content=None):
    """Add pricing page following Page 6 specs."""
    # Section Title - Heading 1
    doc.add_paragraph("PRICING & FINAL NOTES", style='Heading 1')
    
    # Pricing intro
    pricing_intro = doc.add_paragraph()
    pricing_intro.add_run("Our pricing structure is designed to provide maximum value while maintaining competitive rates. All pricing is based on the specific requirements outlined in the RFP document.")
    
    # Pricing table
    doc.add_paragraph("PRICING SCHEDULE", style='Heading 2')
    
    # Create pricing table - Use table
    price_table = doc.add_table(rows=5, cols=2)
    price_table.style = 'Light Grid Accent 1'
    
    # Header row
    header_cells = price_table.rows[0].cells
    header_cells[0].text = "Service Category"
    header_cells[1].text = "Rate"
    
    # Style header row
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            for run in paragraph.runs:
                run.bold = True
    
    # Sample pricing data (would be customized based on actual offerings)
    pricing_data = [
        ("Administrative Staffing", "$XX.XX - $XX.XX per hour"),
        ("Technical Staffing", "$XX.XX - $XX.XX per hour"),
        ("Management Staffing", "$XX.XX - $XX.XX per hour"),
        ("Special Services", "Custom quote based on requirements")
    ]
    
    for i, (service, rate) in enumerate(pricing_data, 1):
        row = price_table.rows[i].cells
        row[0].text = service
        row[1].text = rate
    
    # Optional Services - Italicize
    doc.add_paragraph("OPTIONAL SERVICES", style='Heading 2')
    
    optional_services = [
        "Extended support hours",
        "On-site management",
        "Specialized training",
        "Performance analytics and reporting"
    ]
    
    for service in optional_services:
        bullet = doc.add_paragraph(style='List Bullet')
        bullet_run = bullet.add_run(service)
        bullet_run.italic = True
    
    # Call to Action - Center-aligned paragraph, bold
    cta = doc.add_paragraph()
    cta_run = cta.add_run("CONTACT US TODAY TO DISCUSS YOUR SPECIFIC REQUIREMENTS")
    cta_run.bold = True
    cta.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Signature Line - Use a table for structured signing
    doc.add_paragraph("AUTHORIZATION", style='Heading 2')
    
    # Signature table
    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.cell(0, 0).text = "Authorized Representative:"
    sig_table.cell(0, 1).text = company_data.get('Authorized Representative', '')
    sig_table.cell(1, 0).text = "Title:"
    sig_table.cell(1, 1).text = company_data.get('Authorized Representative Title', '')
    sig_table.cell(2, 0).text = "Signature:"
    sig_table.cell(2, 1).text = "_______________________________"
    
    return doc

def generate_proposal_document(folder_name=None, get_rag_content_func=None, output_path="generated_proposal.docx"):
    """
    Generate a professionally formatted proposal document using:
    1. Company data from CSV
    2. RAG-generated content from processed documents
    
    Following the 6-page structure with specific styling requirements.
    
    Parameters:
    - folder_name: The folder name for the RAG content (passed to get_rag_content_func)
    - get_rag_content_func: A function that takes folder_name and returns a dict with RAG content
    - output_path: Where to save the document
    """
    # Load company data
    company_data = load_company_data()
    
    # Get RAG content if the function is provided
    rag_content = None
    if get_rag_content_func and folder_name:
        rag_content = get_rag_content_func(folder_name)
    
    # Create a new Document
    doc = Document()
    
    # Apply consistent styling
    doc = apply_document_styling(doc)
    
    # Page 1: Cover Page
    doc = add_cover_page(doc, company_data)
    
    # Page 2: Executive Summary
    doc = add_executive_summary(doc, company_data, rag_content)
    
    # Page 3: Company Overview & Experience
    doc = add_company_overview(doc, company_data)
    
    # Page 4: Scope of Work & Solution Approach
    doc = add_scope_and_approach(doc, rag_content)
    
    # Page 5: Key Personnel & Staffing Plan
    doc = add_personnel_section(doc, company_data)
    
    # Page 6: Pricing & Final Notes
    doc = add_pricing_section(doc, company_data, rag_content)
    
    # Save the document
    doc.save(output_path)
    print(f"Document generated successfully: {output_path}")
    
    return output_path

# Example usage:
if __name__ == "__main__":
    # This would be called with the selected document from the Streamlit app
    # folder_name is the ID of the processed document in the FAISS store
    # generate_proposal_document("example_document_abc123")
    pass